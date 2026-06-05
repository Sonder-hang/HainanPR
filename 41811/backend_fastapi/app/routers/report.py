"""报表生成路由（静态简报 + 实时数据报表）"""
import io
import sys
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Any

import httpx
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.database import get_db
from app.models.indicator import Indicator

router = APIRouter(tags=["报表生成"])
logger = logging.getLogger(__name__)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement("w:" + side)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "b8c9e8")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def style_header_cell(cell, text):
    cell.text = text
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_bg(cell, "E8F5E9")
    set_cell_border(cell)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x26, 0x4D)


def style_data_cell(cell, text):
    cell.text = str(text)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x26, 0x4D)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        style_header_cell(table.cell(0, j), h)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            style_data_cell(table.cell(i + 1, j), val)
    return table


def add_rule_section(doc, rule_num, rule_title, basis_label, legal_basis,
                     rule_def, data_source, data_findings,
                     before_table_title, table_headers, table_rows, risk_tip):
    p = doc.add_paragraph()
    run = p.add_run("规则" + rule_num + "（" + rule_title + "）")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x26, 0x4D)

    for label, content in [
        (basis_label, legal_basis),
        ("规则定义", rule_def),
        ("数据来源", data_source),
        ("数据发现", data_findings),
    ]:
        p2 = doc.add_paragraph()
        run2 = p2.add_run(label + "：" + content)
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x1F, 0x26, 0x4D)

    if before_table_title:
        for cap in before_table_title.split("\n"):
            p_cap = doc.add_paragraph()
            run_cap = p_cap.add_run(cap)
            run_cap.font.size = Pt(10)
            run_cap.font.color.rgb = RGBColor(0x1F, 0x26, 0x4D)

    if table_headers and table_rows:
        add_table(doc, table_headers, table_rows)

    if risk_tip:
        p_tip = doc.add_paragraph()
        run_lbl = p_tip.add_run("（五）业务部门风险提示：")
        run_lbl.bold = True
        run_lbl.font.size = Pt(10)
        run_lbl.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
        run_txt = p_tip.add_run(risk_tip)
        run_txt.font.size = Pt(10)
        run_txt.font.color.rgb = RGBColor(0x1F, 0x26, 0x4D)
    doc.add_paragraph()


def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x26, 0x4D)


def generate_brief_report(period_number: str = None, period_year: int = None, period_month: int = None):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    add_para(doc, "内部资料")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("三医智慧监管大数据分析")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x26, 0x4D)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year = period_year if period_year is not None else datetime.now().year
    month = period_month if period_month is not None else datetime.now().month
    pnum = period_number if period_number is not None else "1"
    run = p_sub.add_run("（第" + pnum + "期，" + str(year) + "年" + str(month) + "月）")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x60, 0x80)

    add_para(doc, "按照三医智慧监管工作要求，省卫生健康委医政处、规划信息处，省三医大数据中心利用三医平台数据整合以下资源库并进行分析：一是人员信息库，整合医院HIS系统（含员工、职称数据）与医护注册系统（含医师身份与执业信息），关联生成统一人员档案。二是机构信息库，基于医护注册系统内医疗机构执业许可信息及诊疗科目名录，构建\u201c诊疗科目—科室—诊疗行为\u201d关联规则库。三是技术信息库，融合三医平台病案首页编码与限制类技术备案记录，并依据国家限制类技术目录明确对应手术操作。数据发现如下：")

    # 规则一
    add_rule_section(doc, rule_num="一",
        rule_title="机构要素：医疗机构床位面积比符合性监测",
        basis_label="监管依据",
        legal_basis="卫生部《医疗机构基本标准（试行）》（卫医发〔1994〕第30号）、《医疗机构校验管理办法》。",
        rule_def="核算医疗机构用于住院服务的实际建筑面积与实有床位数是否符合国家分类别设定的最低标准。床位面积比 = 用于住院服务的建筑面积 \u00f7 实有床位数。",
        data_source="医疗卫生机构年报表（卫健统1-表）。",
        data_findings="2024年全年不符合\u201c每床建筑面积\u201d标准的医疗机构共14家。医疗机构计算标准床位数与上报床位数示例如下，明细数据详见附件-1。",
        before_table_title="医疗机构计算标准床位数与上报床位数示例",
        table_headers=["机构名称", "医院等级", "机构类别", "年末房屋建筑面积（平方米）", "每床建筑面积（标准-不少于）", "计算标准床位数", "年报-实有床位"],
        table_rows=[
            ["三亚***", "一级", "综合医院", "858", "45", "19.07", "20"],
            ["万宁***", "一级", "综合医院", "674", "45", "14.98", "70"],
            ["海南医学院***", "二级", "综合医院", "4108", "45", "91.29", "300"],
            ["海口方卓***", "一级", "综合医院", "600", "45", "13.33", "20"],
            ["儋州***", "一级", "综合医院", "3416", "45", "75.91", "96"],
            ["万宁市*****", "二级", "妇幼保健院", "1690", "45", "37.56", "41"],
            ["昌江红*****", "二级", "精神病医院", "5403", "40", "135.08", "255"],
            ["万宁****", "一级", "综合医院", "2100", "45", "46.67", "49"],
            ["博鳌****", "三级", "综合医院", "29953", "60", "499.22", "500"],
            ["乐东宁康***", "二级", "精神病医院", "1", "40", "0.03", "201"],
            ["海南****", "/", "疗养院", "3321", "40", "73.80", "190"],
            ["白沙康民中西医结合****", "一级", "中西医结合医院", "500", "35", "14.29", "60"],
            ["海南省****", "三级", "精神病医院", "22385", "45", "497.44", "1270"],
            ["海南****", "二级", "精神病医院", "6600", "40", "165.00", "299"],
        ],
        risk_tip="上述医疗机构可能存在违反《医疗机构基本标准（试行）》，超标准设置床位，存在医疗质量安全隐患；可能存在医疗机构设置审批、校验等环节纰漏等问题；可能存在违反卫生统计相关法律法规，未及时更新机构相关数据情况。")
    add_para(doc, "属地卫生健康委须现场核实情况，指导医疗机构及时更新数据，按标准核减床位，加强对医疗质量的督导检查。")

    # 规则二
    add_rule_section(doc, rule_num="二",
        rule_title="人员要素：抗菌药物分级管理监测",
        basis_label="监管依据",
        legal_basis="《抗菌药物临床应用管理办法》第六条、第二十四条；《海南省医疗机构抗菌药物临床应用分级管理目录（2024年版）》。",
        rule_def="医师越级开具限制使用级/特殊使用级抗菌药物。含：初级任职资格医师不可授予限制级、特殊使用级抗菌药物的处方权；中级任职资格医师不可授予特殊使用级抗菌药物的处方权。（省卫生健康委医政处）",
        data_source="三医平台住院医嘱数据、门诊处方数据、医院HIS员工（职工）表、职称字典表。",
        data_findings="针对\u201c头孢他啶/阿维巴坦\u201d\u201c亚胺培南/西司他丁\u201d2种特殊使用级抗生素，2025年全年共存在231条\u201c医师开具与职称不符的特殊使用级抗生素\u201d的数据。共涉及87位住院或者主治医师、6家医疗机构。",
        before_table_title="例：医疗机构开具特殊级抗生素统计",
        table_headers=["序号", "医疗机构名称", "住院/主治医师开具特殊级抗生素次数"],
        table_rows=[
            ["1", "海南省***医院", "146"],
            ["2", "海南省***医院", "74"],
            ["3", "海南省***医院", "5"],
            ["4", "海南省***医院", "3"],
            ["5", "海南省***医院", "2"],
            ["6", "海南省***医院", "1"],
        ],
        risk_tip="根据《抗菌药物临床应用管理办法》第二十七条，特殊使用级抗生素须由抗菌药物管理工作组指定的感染性疾病科、呼吸科、重症医学科、微生物检验科、药学部门等具有高级专业技术职务任职的医师、药师会诊；第二十八条，因抢救生命垂危的患者等紧急情况，医师可以越级使用抗菌药物，应当详细记录用药指征，于24小时内补办越级使用抗菌药物的必要手续。")
    add_para(doc, "综上，上述医院要根据我委提供的分析数据，有针对性的对上述数据进行核查，重点查看会诊记录、记录用药指征情况及补办越级使用手续情况。")

    # 规则三
    add_rule_section(doc, rule_num="三",
        rule_title="人员要素：医师跨机构诊疗异常数据监测",
        basis_label="监管依据",
        legal_basis="《医师执业注册管理办法》《关于推进和规范医师多点执业的若干意见》《海南省卫生健康委关于加强公立医疗机构医师多点执业管理的通知》。",
        rule_def="同一医师在相近时间内出现在两家及以上医疗机构开具医嘱或处方。（省卫生健康委医政处）",
        data_source="三医平台住院医嘱数据、门诊处方数据。",
        data_findings="2025年全年同一医师1分钟内在不同机构同时产生诊疗记录，有23条记录；1—5分钟的有38条记录。共涉及9位医师、4家医疗机构。",
        before_table_title="医师诊疗活动情况示例",
        table_headers=["开单医师姓名", "诊疗时间差(分钟)", "本次开单机构", "本次开单时间", "本次诊疗环节", "上次开单机构", "上次开单时间", "上次诊疗环节"],
        table_rows=[
            ["曾**", "5", "海南***", "2025/12/5 10:43:00", "住院", "海南***", "2025/12/5 10:37:25", "门急诊"],
            ["赖**", "3", "海南***", "2025/2/24 9:16:55", "住院", "海南***", "2025/2/24 9:13:45", "门急诊"],
            ["李**", "1", "海南***", "2025/3/20 8:51:37", "住院", "海南***", "2025/3/20 8:50:32", "门急诊"],
            ["卢**", "5", "海南***", "2025/12/1 12:20:00", "住院", "海南***", "2025/12/1 12:14:03", "门急诊"],
            ["石**", "1", "海南***", "2025/11/20 16:32:00", "住院", "海南***", "2025/11/20 16:30:24", "门急诊"],
            ["伟*", "4", "海南***", "2025/12/11 9:18:00", "住院", "海南***", "2025/12/11 9:13:19", "门急诊"],
            ["吴小秦", "0", "海南***", "2025/4/8 15:14:59", "住院", "海南***", "2025/4/8 15:14:17", "门急诊"],
            ["谢*", "0", "海南***", "2025/4/29 11:23:17", "住院", "海南***", "2025/4/29 11:23:12", "门急诊"],
            ["张永*", "2", "海南***", "2025/11/14 9:19:24", "门急诊", "海南***", "2025/11/14 9:17:22", "门急诊"],
        ],
        risk_tip="上述数据说明医疗机构可能存在医师挂证或医师冒名开具处方的问题，上述医疗机构要按照分析情况查阅医师排班表等进一步核实，对于挂证或冒名等问题进一步处理。")

    # 规则四
    add_rule_section(doc, rule_num="四",
        rule_title="人员要素：村卫生室年度服务量低线预警",
        basis_label="监管依据",
        legal_basis="《乡镇卫生院/社区卫生服务中心服务能力评价指南（2023版）》、《全国卫生资源与医疗服务统计调查制度》。",
        rule_def="公立基层医疗机构（村卫生室），一个自然年度内，年收治病人数小于365人次。",
        data_source="医疗卫生机构年报表（卫健统1-3表）村卫生室。",
        data_findings="2024年全年村卫生室接诊人次数少于365人次共753家，其中接诊0人共38家、接诊1-99人共315家、接诊100-299人共353家、接诊300人及以上共47家。",
        before_table_title="",
        table_headers=["行政区划", "年门诊量为0的村卫生室(所)数量"],
        table_rows=[
            ["儋州市", "5"], ["文昌市", "1"], ["万宁市", "8"], ["屯昌县", "6"],
            ["澄迈县", "9"], ["临高县", "1"], ["白沙黎族自治县", "3"],
            ["昌江黎族自治县", "2"], ["乐东黎族自治县", "2"], ["陵水黎族自治县", "1"],
        ],
        risk_tip="")

    # 规则五
    add_rule_section(doc, rule_num="五",
        rule_title="技术要素：单一诊疗项目集中度监测",
        basis_label="监测依据",
        legal_basis="《医疗机构校验管理办法》、《医疗机构诊疗科目名录》。",
        rule_def="连续三个月内单一医疗机构单一诊断和单一术式占比超过全院80%。（省卫生健康委医政处）",
        data_source="医疗卫生机构月报表（卫健统4-1表）。",
        data_findings="2024年全年，单一诊断连续3个月占统计全部诊断比例80%的医院共10家；单一手术连续3个月占统计全部诊断比例80%的医院共4家。",
        before_table_title="医疗机构手术统计情况示例\n医疗机构诊断统计情况示例",
        table_headers=["医疗机构名称", "手术名称", "报表期", "当月出院总人数", "当月当前手术总人数", "占比"],
        table_rows=[
            ["博鳌怡***", "人工耳蜗置入术", "202507", "8", "7", "87.50%"],
            ["博鳌怡***", "人工耳蜗置入术", "202508", "11", "11", "100.00%"],
            ["博鳌怡***", "人工耳蜗置入术", "202509", "12", "12", "100.00%"],
            ["博鳌怡***", "人工耳蜗置入术", "202510", "5", "5", "100.00%"],
            ["博鳌怡***", "人工耳蜗置入术", "202511", "15", "14", "93.33%"],
            ["儋州爱尔**望眼科医院", "白内障超声乳化抽吸术", "202503", "138", "116", "84.06%"],
            ["儋州爱尔**望眼科医院", "白内障超声乳化抽吸术", "202504", "150", "134", "89.33%"],
            ["儋州爱尔**望眼科医院", "白内障超声乳化抽吸术", "202505", "127", "105", "82.68%"],
            ["儋州爱尔**望眼科医院", "白内障超声乳化抽吸术", "202506", "96", "77", "80.21%"],
            ["儋州爱尔**望眼科医院", "白内障超声乳化抽吸术", "202507", "209", "169", "80.86%"],
            ["儋州爱尔**望眼科医院", "白内障超声乳化抽吸术", "202508", "156", "127", "81.41%"],
            ["海口**医院", "内镜下结肠息肉切除术", "202510", "25", "21", "84.00%"],
            ["海口**医院", "内镜下结肠息肉切除术", "202511", "25", "20", "80.00%"],
            ["海口**医院", "内镜下结肠息肉切除术", "202512", "9", "8", "88.89%"],
            ["海口中山医院***", "白内障超声乳化抽吸术", "202503", "191", "155", "81.15%"],
            ["海口中山医院***", "白内障超声乳化抽吸术", "202504", "92", "77", "83.70%"],
            ["海口中山医院***", "白内障超声乳化抽吸术", "202505", "74", "63", "85.14%"],
            ["海口中山医院***", "白内障超声乳化抽吸术", "202506", "77", "66", "85.71%"],
        ],
        risk_tip="上述数据说明医疗机构可能存在诱导住院或过度治疗的风险，属地卫生健康委须进一步现场核实情况，查明患者来源是否过于集中在某地或患者来就诊时是否被宣传诱导，查明患者接受的手术或治疗是否存在明显的指征依据。")

    # 规则六
    add_rule_section(doc, rule_num="六",
        rule_title="医保基金安全：超频次收费监测",
        basis_label="监管依据",
        legal_basis="《医疗保障基金使用监督管理条例》、《海南省医疗服务价格》2024年版。",
        rule_def="监测医师开具的诊疗服务情况，如诊疗服务项目开具频次、数量等。",
        data_source="三医数据平台病案首页、三医数据平台财务收费记录。",
        data_findings="2025年全年部分医疗机构的部分开具了血氧饱和度监测和遥测心电监护项目的病例。筛选出，单日开具某以\u201c小时\u201d为计价单位的项目超过24小时共20条。",
        before_table_title="医疗机构开具项目信息示例",
        table_headers=["机构名称", "住院号", "项目名称", "费用日期", "开具总数量", "开具金额", "判断说明"],
        table_rows=[
            ["海南**医院", "ZY430000149608", "血氧饱和度监测", "2025年8月18日", "30", "165", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "ZY430000131346", "血氧饱和度监测", "2025年8月20日", "26", "143", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "ZY360000502521", "血氧饱和度监测", "2025年8月17日", "58", "319", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "ZY230000190362", "血氧饱和度监测", "2025年8月18日", "27", "148.5", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "25008755", "血氧饱和度监测", "2025年8月28日", "27", "148.5", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "25008342", "血氧饱和度监测", "2025年8月16日", "32", "176", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "25007678", "血氧饱和度监测", "2025年7月31日", "44", "242", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "25007557", "血氧饱和度监测", "2025年8月1日", "31", "170.5", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "2025016712001#1", "血氧饱和度监测", "2025年8月25日", "25.5", "140.25", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "2025016686001#1", "血氧饱和度监测", "2025年8月25日", "26", "143", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "2025016605001#1", "血氧饱和度监测", "2025年8月24日", "32", "176", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "2025016507001#1", "血氧饱和度监测", "2025年8月22日", "29", "159.5", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "2025016339001#1", "血氧饱和度监测", "2025年8月24日", "48", "264", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "90200965_3", "血氧饱和度监测", "2025年8月18日", "32", "176", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "03733116_1", "血氧饱和度监测", "2025年7月25日", "25", "137.5", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "03722046_2", "血氧饱和度监测", "2025年7月9日", "25", "137.5", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "03722046_2", "血氧饱和度监测", "2025年6月27日", "25", "137.5", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "02009423_6", "血氧饱和度监测", "2025年7月29日", "25", "137.5", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "01120719_3", "血氧饱和度监测", "2025年7月25日", "25", "137.5", "血氧饱和度监测不能超过24小时/天"],
            ["海南**医院", "00488036_5", "遥测心电监护", "2025年7月31日", "25", "225", "遥测心电监护不能超过24小时/天"],
        ],
        risk_tip="")

    # 规则七
    add_rule_section(doc, rule_num="七",
        rule_title="特殊群体：未成年人异常诊疗情形监测",
        basis_label="监管依据",
        legal_basis="《关于建立侵害未成年人案件强制报告制度的意见（试行）》、《中华人民共和国未成年人保护法》、《中华人民共和国反家庭暴力法》。",
        rule_def="医疗机构接诊的未成年人异常诊疗情形进行识别，及时发现疑似侵害未成年人线索。识别对象限定：以就诊患者年龄\uff1c18周岁为筛选条件，纳入门诊及住院就诊记录（2）高风险诊断匹配：将患者就诊主诊断 ICD10 编码与\u201c侵害未成年人高风险诊断编码库\u201d（性病、生殖器/泌尿道损伤、流产及妊娠等方向）进行匹配，判定是否命中疑似侵害线索。",
        data_source="三医数据平台病案首页数据，三医数据平台门诊就诊数据，ICD10诊断编码库。",
        data_findings="2025年全年，门诊环节年龄\uff1c18周岁，诊断为流产、梅毒、妊娠的共52条，其中14岁及以下共9条，14岁以上共43条；",
        before_table_title="例：部分医疗机构14岁以下（流产、梅毒、妊娠）住院患者人数统计",
        table_headers=["序号", "机构名称", "住院人数"],
        table_rows=[
            ["1", "乐东***民医院", "14"],
            ["2", "儋州市***医院", "13"],
            ["3", "海南省***医院（秀英院区）", "10"],
            ["4", "海南医学院***医院", "10"],
            ["5", "海南***医院", "9"],
        ],
        risk_tip="上述数据说明医疗机构接诊的未成年患者可能收到不法侵害，属地卫生健康委须指导、提醒、检查医疗机构是否落实强制报告制度。")
    add_para(doc, "住院环节符合年龄\uff1c18周岁，且诊断为流产、梅毒、妊娠的共698条，其中14岁及以下共121条，14岁以上共577条。")

    # 附件清单
    p_at = doc.add_paragraph()
    run = p_at.add_run("附件清单")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x26, 0x4D)

    for a in [
        "     附件 1：医疗机构床位面积比符合性监测",
        "     附件 2：抗菌药物分级管理监测",
        "     附件 3：医师跨机构诊疗异常数据监测",
        "     附件 4：村卫生室年度服务量低线预警",
        "     附件 5：单一诊疗项目集中度监测",
        "     附件 6：超频次收费监测",
        "     附件 7：未成年人异常诊疗情形监测",
    ]:
        add_para(doc, a)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@router.get("/brief-report/download")
def download_brief_report(
    period_number: str = "1",
    period_year: int = None,
    period_month: int = None,
):
    if period_year is None:
        period_year = datetime.now().year
    if period_month is None:
        period_month = datetime.now().month
    buffer = generate_brief_report(period_number, period_year, period_month)
    filename = "三医智慧监管大数据分析简报_第" + period_number + "期_" + str(period_year) + "年" + str(period_month) + "月.docx"
    import urllib.parse
    encoded_filename = urllib.parse.quote(filename)
    cd = "attachment; filename*=UTF-8''" + encoded_filename
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": cd},
    )


# ─────────────────────────────────────────────────────────────────────────────
# 实时数据报表生成
# ─────────────────────────────────────────────────────────────────────────────

class ReportGenerateRequest(BaseModel):
    report_type: str                          # 报表类型 ID，如 "minorProtection"
    time_range: Optional[str] = None           # "month" | "quarter" | "year" | "custom"（兼容旧格式）
    time_mode: Optional[str] = None            # "immediate" | "monthly" | "quarterly"（与指标执行一致）
    time_value: Optional[str] = None           # 如 "2026-05" 或 "2026-Q1"
    start_date: Optional[str] = None           # custom 模式起始日期 "YYYY-MM-DD"
    end_date: Optional[str] = None            # custom 模式结束日期 "YYYY-MM-DD"
    hospital_codes: Optional[List[str]] = None  # 空/null = 全省
    group_by_hospital: bool = True             # 是否按医院分组


def _time_mode_from_range(time_range: str) -> Optional[str]:
    return {"month": "monthly", "quarter": "quarterly", "year": "year"}.get(time_range)


def _time_value_from_range(
    time_range: str,
    time_value: Optional[str],
    start_date: Optional[str],
    end_date: Optional[str],
) -> str:
    if time_range == "custom" and start_date:
        return start_date[:7]  # "YYYY-MM"
    if time_value:
        return time_value
    now = datetime.now()
    if time_range == "month":
        return now.strftime("%Y-%m")
    if time_range == "quarter":
        q = (now.month - 1) // 3 + 1
        return f"{now.year}-Q{q}"
    if time_range == "year":
        return str(now.year)
    return ""


class ReportGenerateResponse(BaseModel):
    ok: bool
    report_type: str
    headers: List[str]
    rows: List[List[Any]]
    total_count: int
    summary: dict
    preview_columns: List[str]
    preview_rows: List[dict]
    message: Optional[str] = None


@router.post("/generate", response_model=ReportGenerateResponse)
def generate_report(data: ReportGenerateRequest, db: Session = Depends(get_db)):
    """
    生成实时数据报表。

    支持真实数据：
    - antibioticManagement      → id=63  抗菌药物分级管理监测（越权开具抗生素）
    - crossInstitutionDiagnosis → id=2   医师跨机构诊疗异常监测（时空轨迹异常）
    - restrictedTechUsage       → id=7   国家级限制性技术使用监测
    - practiceOverdue          → id=3   医师执业超期异常监控
    - practiceLocation         → id=64  医师执业地点异常监控
    - minorProtection          → id=10  未成年人异常诊疗情形监测
    """
    # ── 1. 查指标配置 ──────────────────────────────────────────────────────
    indicator_id_map = {
        "antibioticManagement":      63,
        "crossInstitutionDiagnosis": 2,
        "restrictedTechUsage":       7,
        "practiceOverdue":           3,
        "practiceLocation":         64,
        "minorProtection":           10,
    }

    indicator_id = indicator_id_map.get(data.report_type)

    if indicator_id is None:
        raise HTTPException(status_code=400, detail=f"未知报表类型: {data.report_type}")

    ind = db.query(Indicator).filter(Indicator.id == indicator_id).first()
    if not ind:
        raise HTTPException(status_code=404, detail=f"指标 id={indicator_id} 不存在")

    ind_dict = {
        "id": ind.id,
        "name": ind.name,
        "calc_type": ind.calc_type,
        "sql_content": ind.sql_content,
        "indicator_type": ind.indicator_type,
        "date_field": ind.date_field or "VST_DT_TM",
        "numerator_sql": ind.numerator_sql,
        "denominator_sql": ind.denominator_sql,
        "subitem_config": ind.subitem_config,
    }

    # ── 调试日志：打印 SQL 内容 ─────────────────────────────────────────────
    logger.info(f"[报表生成] 指标ID={ind.id}，名称={ind.name}")
    num_sql_raw = ind.numerator_sql or ""
    den_sql_raw = ind.denominator_sql or ""
    logger.info(f"[报表生成] 分子SQL前100字符: {num_sql_raw[:100]!r}")
    logger.info(f"[报表生成] 分母SQL前100字符: {den_sql_raw[:100]!r}")
    logger.info(f"[报表生成] 分子SQL含反斜杠n数量: {num_sql_raw.count(chr(92)+'n')}")
    logger.info(f"[报表生成] 分母SQL含反斜杠n数量: {den_sql_raw.count(chr(92)+'n')}")

    # ── 2. 构建执行参数 ──────────────────────────────────────────────────────
    ind_dict["hospital_codes"] = data.hospital_codes
    ind_dict["group_by_hospital"] = data.group_by_hospital
    # 优先使用 time_mode（与指标执行一致），否则兼容 time_range
    time_mode = data.time_mode or _time_mode_from_range(data.time_range or "")
    ind_dict["time_mode"] = time_mode
    # time_value: 优先用前端传入的，否则按 time_mode 推算
    if data.time_value:
        ind_dict["time_value"] = data.time_value
    elif time_mode == "monthly":
        now = datetime.now()
        ind_dict["time_value"] = now.strftime("%Y-%m")
    elif time_mode == "quarterly":
        now = datetime.now()
        q = (now.month - 1) // 3 + 1
        ind_dict["time_value"] = f"{now.year}-Q{q}"
    else:
        ind_dict["time_value"] = None

    # ── 3. 调用 Text2SQLService 执行 ─────────────────────────────────────────
    try:
        from app.services.text2sql import Text2SQLService
        service = Text2SQLService()
        result = service.execute_indicator(ind_dict, db_session=db, skip_save=True)
    except Exception as e:
        logger.exception(f"[报表生成] 执行指标 {indicator_id} 失败")
        raise HTTPException(status_code=500, detail=f"指标执行失败: {str(e)}")

    # ── 4. 解析结果 ───────────────────────────────────────────────────────────
    preview_data = result.get("preview_data") or {}
    cols: List[str] = preview_data.get("columns") or result.get("preview_columns") or []
    rows: List[dict] = preview_data.get("rows") or result.get("preview_rows") or []

    # 转换 rows 为普通列表（兼容前端）
    plain_rows: List[List[Any]] = []
    for r in rows:
        if isinstance(r, dict):
            plain_rows.append([r.get(c) for c in cols])
        elif hasattr(r, "__iter__"):
            plain_rows.append(list(r))

    total_count = result.get("count") or result.get("numerator_count") or 0
    if isinstance(total_count, dict):
        total_count = sum(v for v in total_count.values()) if total_count else 0

    # 汇总行
    summary = {}
    if data.group_by_hospital:
        hosp_results = result.get("hospital_results") or []
        summary = {
            "监测机构数": len(hosp_results),
            "预警总数": total_count,
            "执行状态": "成功" if result.get("ok") else "部分失败",
        }
    else:
        summary = {
            "预警总数": total_count,
            "执行状态": "成功" if result.get("ok") else "失败",
        }

    return ReportGenerateResponse(
        ok=result.get("ok", True),
        report_type=data.report_type,
        headers=cols,
        rows=plain_rows,
        total_count=total_count,
        summary=summary,
        preview_columns=cols,
        preview_rows=rows,
        message=result.get("error"),
    )


# ─────────────────────────────────────────────────────────────────────────────
# 完整数据导出（绕过预览 limit，取全量数据）
# ─────────────────────────────────────────────────────────────────────────────

@router.post("/export-full")
def export_report_full(data: ReportGenerateRequest):
    """
    导出报表全量数据（不限条数，用于 Excel 完整下载）。
    与 /generate 的区别在于：始终取全量数据，不做 limit 100 截断。
    """
    # 复用 generate 的指标映射和配置逻辑
    indicator_id_map = {
        "antibioticManagement":      63,
        "crossInstitutionDiagnosis": 2,
        "restrictedTechUsage":      7,
        "practiceOverdue":          3,
        "practiceLocation":        64,
        "minorProtection":         10,
    }
    indicator_id = indicator_id_map.get(data.report_type)
    if indicator_id is None:
        raise HTTPException(status_code=400, detail=f"未知报表类型: {data.report_type}")

    # 直接查库获取指标配置（不走 db session dependency 以简化）
    import pymysql, json as _json
    conn = pymysql.connect(
        host="127.0.0.1", port=3306,
        user="root", password="123456",
        database="hainan_41811", charset="utf8mb4"
    )
    cur = conn.cursor(pymysql.cursors.DictCursor)
    cur.execute(
        "SELECT * FROM indicator WHERE id = %s", (indicator_id,)
    )
    ind = cur.fetchone()
    conn.close()
    if not ind:
        raise HTTPException(status_code=404, detail=f"指标 {indicator_id} 不存在")

    ind_dict = {
        "id": ind["id"],
        "name": ind["name"],
        "calc_type": ind["calc_type"],
        "sql_content": ind["sql_content"],
        "indicator_type": ind["indicator_type"],
        "date_field": ind["date_field"] or "VST_DT_TM",
        "numerator_date_field": ind["numerator_date_field"] or "discharge",
        "denominator_date_field": ind["denominator_date_field"] or "discharge",
        "numerator_sql": ind["numerator_sql"],
        "denominator_sql": ind["denominator_sql"],
        "subitem_config": (
            _json.loads(ind["subitem_config"])
            if ind.get("subitem_config") and ind["subitem_config"].strip()
            else None
        ),
    }

    # 时间参数
    time_mode = data.time_mode or _time_mode_from_range(data.time_range or "")
    ind_dict["time_mode"] = time_mode
    if data.time_value:
        ind_dict["time_value"] = data.time_value
    elif time_mode == "monthly":
        now = datetime.now()
        ind_dict["time_value"] = now.strftime("%Y-%m")
    elif time_mode == "quarterly":
        now = datetime.now()
        q = (now.month - 1) // 3 + 1
        ind_dict["time_value"] = f"{now.year}-Q{q}"
    else:
        ind_dict["time_value"] = None

    ind_dict["hospital_codes"] = data.hospital_codes
    ind_dict["group_by_hospital"] = data.group_by_hospital

    # 直接执行 SQL（不通过 Celery），绕过 limit
    from app.services.text2sql import Text2SQLService
    service = Text2SQLService()

    # 注入过滤
    raw_sql = ind_dict.get("numerator_sql") or ind_dict.get("sql_content") or ""
    if not raw_sql:
        return {"headers": [], "rows": [], "total": 0}

    raw_date_field = ind_dict.get("numerator_date_field") or ind_dict.get("date_field") or "discharge"
    time_col = service._resolve_time_col(raw_date_field)
    injected = service._inject_filters(
        raw_sql,
        hospital_codes=data.hospital_codes,
        time_mode=time_mode,
        time_value=ind_dict.get("time_value"),
        time_col_name=time_col,
        is_aggregate=False,
    )

    # 全量执行
    cnt, err, cols, rows = service._exec_sql(injected, include_rows=True, fetch_all=True)
    if err:
        raise HTTPException(status_code=500, detail=f"SQL执行失败: {err}")

    # 转为列表格式
    plain_rows_out = []
    for r in rows:
        if isinstance(r, dict):
            plain_rows_out.append([r.get(c) for c in cols])
        else:
            plain_rows_out.append(list(r))

    return {
        "headers": cols,
        "rows": plain_rows_out,
        "total": len(plain_rows_out),
    }
